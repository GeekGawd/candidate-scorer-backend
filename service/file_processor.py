import base64
import io
from typing import Tuple, Optional
import logging
from pypdf import PdfReader
from docx import Document
import re
from fastapi import UploadFile

logger = logging.getLogger(__name__)

class FileProcessor:
    """Service for processing uploaded resume files"""
    
    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file with improved error handling"""
        try:
            pdf_stream = io.BytesIO(file_content)
            reader = PdfReader(pdf_stream)
            
            # Check if PDF has pages
            if len(reader.pages) == 0:
                raise ValueError("PDF file contains no pages")
            
            text = ""
            pages_processed = 0
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        pages_processed += 1
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if pages_processed == 0:
                raise ValueError("No text could be extracted from any pages in the PDF")
            
            extracted_text = text.strip()
            if not extracted_text:
                raise ValueError("PDF appears to be empty or contains only images")
            
            logger.info(f"Successfully extracted text from {pages_processed}/{len(reader.pages)} pages")
            return extracted_text
            
        except ValueError:
            # Re-raise ValueError as-is
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise ValueError("Failed to process PDF file. The file may be corrupted or password-protected.")
    
    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from Word document with improved error handling"""
        try:
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            
            text = ""
            paragraphs_processed = 0
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraphs_processed += 1
            
            # Also extract text from tables
            tables_processed = 0
            for table in doc.tables:
                try:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                text += cell.text + "\n"
                    tables_processed += 1
                except Exception as e:
                    logger.warning(f"Failed to extract text from table: {e}")
                    continue
            
            extracted_text = text.strip()
            if not extracted_text:
                raise ValueError("DOCX file appears to be empty")
            
            logger.info(f"Successfully extracted text from {paragraphs_processed} paragraphs and {tables_processed} tables")
            return extracted_text
            
        except ValueError:
            # Re-raise ValueError as-is
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError("Failed to process Word document. The file may be corrupted or in an unsupported format.")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\-\.\,\(\)\@\:\;]', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    @classmethod
    async def process_resume_file(cls, upload_file: UploadFile) -> Tuple[str, bytes]:
        """
        Process uploaded resume file and extract text
        
        Args:
            upload_file: FastAPI UploadFile object
        
        Returns:
            Tuple of (extracted_text, original_file_bytes)
        """
        filename = getattr(upload_file, 'filename', None)
        
        # Validate file type
        if not filename:
            raise ValueError("No filename provided")
        
        file_extension = filename.lower().split('.')[-1]
        if file_extension not in ['pdf', 'docx']:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
        
        # Read file content with better error handling
        file_content = None
        try:
            # Reset file pointer to beginning in case it was read before
            await upload_file.seek(0)
            file_content = await upload_file.read()
            if not file_content:
                raise ValueError("Empty file uploaded")
        except Exception as e:
            logger.error(f"Failed to read uploaded file {filename}: {str(e)}")
            raise ValueError("Failed to read uploaded file")
        
        # Validate file size
        file_size = len(file_content)
        max_size = 10 * 1024 * 1024  # 10MB
        if file_size > max_size:
            raise ValueError(f"File size {file_size} bytes exceeds maximum size of {max_size} bytes")
        
        # Validate file content for basic structure
        try:
            if file_extension == 'pdf':
                if not file_content.startswith(b'%PDF'):
                    raise ValueError("Invalid PDF file format")
            elif file_extension == 'docx':
                if not file_content.startswith(b'PK'):
                    raise ValueError("Invalid DOCX file format")
        except Exception as e:
            raise ValueError(f"File validation failed: {str(e)}")
        
        # Extract text based on file type
        extracted_text = ""
        try:
            if file_extension == 'pdf':
                extracted_text = await cls.extract_text_from_pdf(file_content)
            elif file_extension == 'docx':
                extracted_text = await cls.extract_text_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_extension} file {filename}: {str(e)}")
            raise ValueError(f"Failed to process {file_extension} file. Please ensure the file is not corrupted.")
        
        # Clean the extracted text
        cleaned_text = cls.clean_text(extracted_text)
        
        if not cleaned_text or len(cleaned_text.strip()) < 50:
            raise ValueError("No meaningful text could be extracted from the file. Please ensure the file contains readable text.")
        
        logger.info(f"Successfully extracted {len(cleaned_text)} characters from {file_extension} file: {filename}")
        
        return cleaned_text, file_content
    
    @staticmethod
    def validate_file_size(upload_file: UploadFile, max_size_mb: int = 10) -> bool:
        """Validate file size (this needs to be called before reading the file)"""
        if hasattr(upload_file, 'size') and upload_file.size:
            max_size_bytes = max_size_mb * 1024 * 1024
            if upload_file.size > max_size_bytes:
                raise ValueError(f"File size exceeds {max_size_mb}MB limit")
        return True 